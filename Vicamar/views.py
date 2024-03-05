from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from docx import Document
from django.views.decorators.csrf import csrf_exempt
from  .forms import CotizacionForm
from docx2pdf import convert
import os
from datetime import datetime

def actualizar_datos(doc, datos:dict[str]):
    for header in datos.keys():
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if f"<<{header}>>" in run.text:
                                run.text = str(datos[header]) 

                                

# Create your views here.
def formularioCotizacion(request):
    if request.method == "POST":
        form = CotizacionForm(request.POST)
        if form.is_valid():
            nombre_doc = "Vicamar/formato.docx"
            doc = Document(nombre_doc)
            form = form.cleaned_data
            data = {}
            for campo in form.keys():
                data[campo.upper()] = form[campo]
            data["FECHA"] = datetime.now().strftime("%d/%m/%Y")
            actualizar_datos(doc, data)    
            doc.save("formato.docx")
            convert("formato.docx", "COTIZACION DE VICAMAR.pdf")
            os.remove("formato.docx")
            form = CotizacionForm()
            with open("COTIZACION DE VICAMAR.pdf", "rb") as pdf_file:
                response = HttpResponse(pdf_file.read(), content_type='application/pdf')
                response['Content-Disposition'] = 'attachment; filename="COTIZACION DE VICAMAR.pdf"'

            os.remove("COTIZACION DE VICAMAR.pdf")
            return response
                        
    else:
        form = CotizacionForm()
    return render(request, "Vicamar/cotizaciones.html", {"form": form})
    
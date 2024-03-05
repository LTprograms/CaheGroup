from typing import Any
from django.db import models
from docx2pdf import convert
import os
from datetime import datetime
from docx import Document
from django.http import HttpResponse
from django.utils.safestring import mark_safe
from django.core.files import File
from django.utils import timezone


# Create your models here.
class Cliente(models.Model):
    ruc = models.CharField(max_length = 15)
    telefono = models.CharField(max_length = 15)
    correo = models.EmailField()
    razon_social = models.CharField(max_length = 30)

    def __str__(self) -> str:
        return self.razon_social
    
    class Meta:
        verbose_name_plural = "Clientes"

class Maquinaria(models.Model):
    nombre = models.CharField(max_length = 30)
    stock = models.IntegerField()
    precio = models.FloatField()
    modelo = models.CharField(max_length = 30)
    marca = models.CharField(max_length = 30)
    activo = models.BooleanField(default=True)
    def __str__(self) -> str:
        return self.nombre
    
    class Meta:
        verbose_name_plural = "Maquinarias"


class Alquiler(models.Model):
    cliente = models.ForeignKey(Cliente, on_delete=models.CASCADE)
    maquinaria = models.ForeignKey(Maquinaria, on_delete=models.CASCADE)
    destino = models.CharField(null=True, blank=True, max_length=30)
    inicio = models.DateField(auto_now_add=True)
    fin = models.DateField(null=True, blank=True)
    cantidad = models.IntegerField(null=True, blank=True)
    descuento = models.FloatField(null=True, blank=True, default=0)
    sub_total = models.FloatField(null=True, blank=True)
    pdf_url = models.FileField(null=True, blank=True, upload_to="cotizaciones/")

    class Meta:
        verbose_name_plural = "Alquileres"  

    def __str__(self) -> str:
        return f"{self.cliente} - {self.maquinaria}"
    
    def pdf(self):
        return mark_safe("<a class='btn btn-success' download='Cotizacion.pdf' href='{url}'>Descargar</a>".format(url=self.pdf_url.url))
    
    def actualizar_datos(doc, datos:dict[str]):
        for header in datos.keys():
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if f"<<{header}>>" in run.text:
                                    run.text = str(datos[header]) 

    def generar_pdf(self, data):
        nombre_doc = "Vicamar/formato.docx"
        doc = Document(nombre_doc)
        subtotal = data['maquinaria'].precio*data['cantidad']*(1 - (data['descuento']/100))
        dias = data['fin'] - timezone.now().date()        
        print(dias)
        dataDic = {
            "CLIENTE" : data['cliente'].razon_social,
            "RUC" : data['cliente'].ruc,
            "DESTINO" : data['destino'],
            "CONTACTO" : data['cliente'].correo,
            "NUMERO_COTIZACION" : self.get_num_cotizacion(),
            "PRECIO" : data['maquinaria'].precio,
            "CANTIDAD" : data['cantidad'],
            "DIAS" : dias.days,
            "DESCUENTO": data['descuento'],
            "SUBTOTAL" : round(subtotal, 2),
            "TOTAL" : round(subtotal*1.18, 2),
        }       

        dataDic["FECHA"] = datetime.now().strftime("%d/%m/%Y")
        Alquiler.actualizar_datos(doc, dataDic)    
        doc.save("formato.docx")
        convert("formato.docx", "COTIZACION DE VICAMAR.pdf")
        os.remove("formato.docx")

    def get_num_cotizacion(self) -> str:
        id = len(Alquiler.objects.all()) + 1
        num = ''
        for i in range(4 - len(str(id))):
            num += '0'
        return num + str(id)

    """ def save(self):
        self.generar_pdf()
        with open("./COTIZACION DE VICAMAR.pdf", "rb") as file:
            self.pdf_url.save("COTIZACION DE VICAMAR.pdf", File(file), save=False)
        os.remove("./COTIZACION DE VICAMAR.pdf") """

    
